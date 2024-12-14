from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import random
import time
from twitter_client import GreenPillBotTester

class BotTestLogger:
    def __init__(self):
        self.doc = Document()
        self.bot = GreenPillBotTester()
        self._setup_document()
        
    def _setup_document(self):
        # Add title
        title = self.doc.add_heading('GreenPillBot Testing Log', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = self.doc.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp.add_run(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        self.doc.add_paragraph()  # Add spacing
        
    def _add_section_header(self, text):
        heading = self.doc.add_heading(text, level=1)
        heading.style.font.size = Pt(14)
        heading.style.font.color.rgb = RGBColor(0, 0, 139)
        
    def _add_tweet(self, tweet, tweet_number):
        p = self.doc.add_paragraph()
        p.add_run(f'Tweet #{tweet_number}\n').bold = True
        p.add_run(tweet)
        p.add_run(f'\nCharacter count: {len(tweet)}')
        self.doc.add_paragraph()  # Add spacing
        
    def _add_reply_test(self, mention, is_relevant, reply=None):
        p = self.doc.add_paragraph()
        p.add_run('Test Mention: ').bold = True
        p.add_run(mention)
        p.add_run('\nRelevance: ').bold = True
        
        if is_relevant:
            p.add_run('Relevant').font.color.rgb = RGBColor(0, 128, 0)
            p.add_run('\nBot Reply: ').bold = True
            p.add_run(reply)
            p.add_run(f'\nCharacter count: {len(reply)}')
        else:
            p.add_run('Not Relevant').font.color.rgb = RGBColor(255, 0, 0)
            
        self.doc.add_paragraph()  # Add spacing
        
    def generate_test_data(self, num_tweets=10, num_replies=10):
        # Generate and log tweets
        self._add_section_header('Generated Tweets')
        for i in range(num_tweets):
            tweet = self.bot.generate_tweet()
            self._add_tweet(tweet, i+1)
            time.sleep(1)  # Add delay to avoid rate limits
            
        # Test replies with mix of relevant and irrelevant mentions
        self._add_section_header('Reply Tests')
        
        # Sample test mentions
        relevant_mentions = [
            "@GreenPillBot loving the work on impact certificates! How do you measure social impact?",
            "@GreenPillBot been diving into mechanism design lately, any recommendations?",
            "@GreenPillBot what's your take on quadratic funding for public goods?",
            "@GreenPillBot how can DAOs improve their governance models?",
            "@GreenPillBot thoughts on retroactive funding for impact?"
        ]
        
        irrelevant_mentions = [
            "@GreenPillBot what's your favorite color?",
            "@GreenPillBot do you like pizza?",
            "@GreenPillBot when is the next bitcoin halving?",
            "@GreenPillBot what's the weather like?",
            "@GreenPillBot who will win the super bowl?"
        ]
        
        # Mix relevant and irrelevant mentions
        all_mentions = relevant_mentions + irrelevant_mentions
        random.shuffle(all_mentions)
        
        for i, mention in enumerate(all_mentions[:num_replies]):
            is_relevant = self.bot.check_relevance(mention)
            reply = self.bot.generate_reply(mention) if is_relevant else None
            self._add_reply_test(mention, is_relevant, reply)
            time.sleep(1)  # Add delay to avoid rate limits
            
    def save_document(self, filename='bot_test_log.docx'):
        # Add summary
        self._add_section_header('Test Summary')
        summary = self.doc.add_paragraph()
        summary.add_run('This document contains test outputs from the GreenPillBot, demonstrating its tweet generation and reply capabilities. The bot uses advanced language models and maintains context awareness while generating responses.')
        
        # Save the document
        self.doc.save(filename)
        print(f"Test log saved to {filename}")

def run_tests():
    logger = BotTestLogger()
    print("Starting bot tests...")
    print("Generating tweets and testing replies...")
    logger.generate_test_data(num_tweets=10, num_replies=10)
    print("Saving test log...")
    logger.save_document()
    print("Testing complete!")

if __name__ == "__main__":
    run_tests()